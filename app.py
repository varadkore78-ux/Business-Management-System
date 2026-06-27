import streamlit as st
import sqlite3
import os
import uuid

from datetime import datetime
from docx import Document
from docx.shared import Inches
import fitz
from PyPDF2 import PdfReader

# ==========================================
# PAGE SETTINGS
# ==========================================

st.set_page_config(
    
    page_title="Surabhi Plastics",
    page_icon="🏭",
    layout="wide"
)

st.title("🏭 SURABHI PLASTICS MANAGEMENT SYSTEM")

# ==========================================
# UPLOAD FOLDER
# ==========================================

UPLOAD_FOLDER = "uploads"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ==========================================
# DATABASE CONNECTION
# ==========================================

conn = sqlite3.connect(
    "business.db",
    check_same_thread=False
)

cursor = conn.cursor()

# ==========================================
# CUSTOMER TABLE
# ==========================================

cursor.execute("""
CREATE TABLE IF NOT EXISTS customers(
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT,
    gst TEXT,
    phone TEXT,
    organization TEXT,
    city TEXT
)
""")

# ==========================================
# ORDER TABLE
# ==========================================

cursor.execute("""
CREATE TABLE IF NOT EXISTS orders(
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    invoice TEXT,
    customer TEXT,
    product TEXT,
    order_details TEXT,
    quantity INTEGER,
    rate REAL,
    total REAL,
    file_path TEXT
)
""")

conn.commit()

# ==========================================
# UPDATE OLD DATABASE
# ==========================================

try:
    cursor.execute("""
    ALTER TABLE orders
    ADD COLUMN file_path TEXT
    """)
    conn.commit()

except sqlite3.OperationalError:
    pass

# ==========================================
# AUTO INVOICE NUMBER
# ==========================================

cursor.execute("SELECT COUNT(*) FROM orders")

count = cursor.fetchone()[0] + 1

invoice = f"INV{count:03d}"

# ==========================================
# ADD CUSTOMER
# ==========================================

st.markdown("---")
st.header("👤 Add Customer")

cust_name = st.text_input("Customer Name")

gst = st.text_input("GST Number")

phone = st.text_input("Phone Number")

organization = st.text_input("Organization")

city = st.text_input("City")

col1, col2 = st.columns(2)

with col1:

    if st.button("➕ Save Customer"):

        if cust_name.strip() == "":

            st.error("Customer Name is required.")

        else:

            cursor.execute("""
            INSERT INTO customers
            (
                name,
                gst,
                phone,
                organization,
                city
            )
            VALUES
            (
                ?,?,?,?,?
            )
            """,
            (
                cust_name,
                gst,
                phone,
                organization,
                city
            ))

            conn.commit()

            st.success("✅ Customer Saved Successfully!")

with col2:

    if st.button("🗑 Clear Customer Form"):

        st.rerun()

# ==========================================
# CUSTOMER LIST
# ==========================================

st.markdown("---")

st.header("📋 Customer List")

cursor.execute("""
SELECT
    id,
    name,
    gst,
    phone,
    organization,
    city
FROM customers
ORDER BY id DESC
""")

customers = cursor.fetchall()

if customers:

    st.dataframe(
        customers,
        use_container_width=True
    )

else:

    st.info("No Customers Available.")

    # ==========================================
# ADD ORDER
# ==========================================

st.markdown("---")
st.header("📦 Add Order")

st.text_input(
    "Invoice Number",
    value=invoice,
    disabled=True
)

customer = st.text_input(
    "Customer Name",
    key="order_customer"
)

product = st.text_input("Product Name")

order_details = st.text_area(
    "Order Details",
    height=120
)

qty = st.number_input(
    "Quantity",
    min_value=1,
    value=1
)

rate = st.number_input(
    "Rate",
    min_value=0.0,
    value=0.0
)

total = qty * rate

st.write(f"### Total Amount : ₹ {total:.2f}")

order_file = st.file_uploader(
    "Attach File",
    type=[
        "doc",
        "docx",
        "pdf",
        "xls",
        "xlsx",
        "jpg",
        "jpeg",
        "png"
    ]
)

if st.button("💾 Save Order"):

    if customer.strip() == "":
        st.error("Please enter Customer Name.")

    elif product.strip() == "":
        st.error("Please enter Product Name.")

    else:

        file_path = ""

        if order_file is not None:

            unique_name = str(uuid.uuid4()) + "_" + order_file.name

            file_path = os.path.join(
                UPLOAD_FOLDER,
                unique_name
            )

            with open(file_path, "wb") as f:
                f.write(order_file.getbuffer())

        cursor.execute("""
        INSERT INTO orders
        (
            invoice,
            customer,
            product,
            order_details,
            quantity,
            rate,
            total,
            file_path
        )
        VALUES
        (
            ?,?,?,?,?,?,?,?
        )
        """,
        (
            invoice,
            customer,
            product,
            order_details,
            qty,
            rate,
            total,
            file_path
        ))

        conn.commit()

        st.success("✅ Order Saved Successfully!")

        # ==========================================
# ORDER LIST
# ==========================================

st.markdown("---")
st.header("📋 Order List")

cursor.execute("""
SELECT
    invoice,
    customer,
    product,
    quantity,
    rate,
    total,
    file_path
FROM orders
ORDER BY id DESC
""")

orders = cursor.fetchall()

if orders:

    table = []

    for row in orders:

        table.append({
            "Invoice": row[0],
            "Customer": row[1],
            "Product": row[2],
            "Quantity": row[3],
            "Rate": row[4],
            "Total": row[5],
            "Attachment": "Yes" if row[6] else "No"
        })

    st.dataframe(
        table,
        use_container_width=True
    )

else:

    st.info("No Orders Available.")

    # ==========================================
# SEARCH CUSTOMER
# ==========================================

st.markdown("---")
st.header("🔍 Search Customer")

search_customer = st.text_input(
    "Enter Customer Name",
    key="search_customer"
)

if st.button("Search Customer"):

    cursor.execute("""
    SELECT
        id,
        name,
        gst,
        phone,
        organization,
        city
    FROM customers
    WHERE name LIKE ?
    """, ('%' + search_customer + '%',))

    result = cursor.fetchall()

    if result:

        st.success("Customer Found")

        st.dataframe(
            result,
            use_container_width=True
        )

    else:

        st.warning("Customer Not Found")

        # ==========================================
# SEARCH ORDERS
# ==========================================

st.markdown("---")
st.header("🔍 Search Orders")

search_order = st.text_input(
    "Enter Customer Name",
    key="search_order"
)

if st.button("Search Orders"):

    cursor.execute("""
    SELECT
        invoice,
        customer,
        product,
        order_details,
        quantity,
        rate,
        total,
        file_path
    FROM orders
    WHERE customer LIKE ?
    ORDER BY id DESC
    """, ('%' + search_order + '%',))

    orders = cursor.fetchall()

    if orders:

        for row in orders:

            st.subheader(f"📄 Invoice : {row[0]}")

            st.write("👤 Customer :", row[1])
            st.write("📦 Product :", row[2])
            st.write("📝 Order Details :", row[3])
            st.write("🔢 Quantity :", row[4])
            st.write("💰 Rate : ₹", row[5])
            st.write(f"💵 Total : ₹ {row[6]:.2f}")

            if row[7]:

                if os.path.exists(row[7]):

                    col1, col2 = st.columns(2)

                    with col1:

                        if st.button(
                            "📂 Open Attachment",
                            key=f"open_{row[0]}"
                        ):
                            os.startfile(row[7])

                    with col2:

                        with open(row[7], "rb") as f:

                            st.download_button(
                                "⬇ Download Attachment",
                                data=f,
                                file_name=os.path.basename(row[7]),
                                mime="application/octet-stream",
                                key=f"download_{row[0]}"
                            )

                else:

                    st.error("Attachment file not found.")

            else:

                st.info("No attachment uploaded.")

            st.divider()

    else:

        st.warning("No Orders Found.")



# ==========================================
# GENERATE SINGLE INVOICE
# ==========================================

st.markdown("---")
st.header("📄 Generate Invoice")

invoice_search = st.text_input(
    "Enter Invoice Number",
    placeholder="Example : INV001"
)

if st.button("Generate Invoice"):

    cursor.execute("""
        SELECT
            invoice,
            customer,
            product,
            order_details,
            quantity,
            rate,
            total,
            file_path
        FROM orders
        WHERE invoice = ?
    """, (invoice_search,))

    row = cursor.fetchone()

    if row is None:

        st.error("Invoice not found.")

    else:

        doc = Document()

        # ---------------- Company Logo ----------------

        if os.path.exists("surabhi.png"):

            doc.add_picture(
                "surabhi.png",
                width=Inches(1.5)
            )

        # ---------------- Company Details ----------------

        doc.add_heading("SURABHI PLASTICS", 0)

        doc.add_paragraph("Kolhapur, Maharashtra")

        doc.add_paragraph(
            "Phone : +91 9960133278, 7875707911"
        )

        doc.add_paragraph(
            "Email : surabhiplastics1994@gmail.com"
        )

        doc.add_paragraph(
            "GST No : 27AFMPK1831L1Z7"
        )

        doc.add_paragraph(
            f"Date : {datetime.now().strftime('%d-%m-%Y')}"
        )

        doc.add_heading(
            "TAX INVOICE",
            level=1
        )

        table = doc.add_table(
            rows=7,
            cols=2
        )

        table.style = "Table Grid"

        table.cell(0,0).text = "Invoice No"
        table.cell(0,1).text = row[0]

        table.cell(1,0).text = "Customer"
        table.cell(1,1).text = row[1]

        table.cell(2,0).text = "Product"
        table.cell(2,1).text = row[2]

        table.cell(3,0).text = "Order Details"
        table.cell(3,1).text = row[3]

        table.cell(4,0).text = "Quantity"
        table.cell(4,1).text = str(row[4])

        table.cell(5,0).text = "Rate"
        table.cell(5,1).text = f"₹ {row[5]}"

        table.cell(6,0).text = "Total"
        table.cell(6,1).text = f"₹ {row[6]}"

        file_path = row[7]

                # ---------------- Attached File ----------------

        if file_path:

            doc.add_page_break()

            doc.add_heading("Attached File", level=1)

            if os.path.exists(file_path):

                extension = os.path.splitext(file_path)[1].lower()

                # -------- Word File --------
                if extension == ".docx":

                    try:

                        attached_doc = Document(file_path)

                        for para in attached_doc.paragraphs:

                            if para.text.strip():

                                doc.add_paragraph(para.text)

                    except Exception as e:

                        doc.add_paragraph(
                            f"Cannot read Word file: {e}"
                        )

                # -------- PDF File --------
                elif extension == ".pdf":

                    try:

                        pdf = fitz.open(file_path)

                        for page_number, page in enumerate(pdf):

                            image_name = f"temp_page_{page_number}.png"

                            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))

                            pix.save(image_name)

                            doc.add_picture(
                                image_name,
                                width=Inches(5.8)
                            )

                        pdf.close()

                    except Exception as e:

                        doc.add_paragraph(
                            f"Cannot read PDF: {e}"
                        )

                # -------- Image File --------
                elif extension in [".jpg", ".jpeg", ".png"]:

                    try:

                        doc.add_picture(
                            file_path,
                            width=Inches(5.8)
                        )

                    except Exception as e:

                        doc.add_paragraph(
                            f"Cannot read Image: {e}"
                        )

                else:

                    doc.add_paragraph(
                        "Preview not available for this file type."
                    )

            else:

                doc.add_paragraph("Attached file not found.")

        # ---------------- Signature ----------------

        doc.add_page_break()

        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("____________________________")
        doc.add_paragraph("Authorized Signature")
        doc.add_paragraph("Surabhi Plastics")

        doc.add_paragraph("")
        doc.add_paragraph("Thank You For Your Business.")
        doc.add_paragraph("Visit Again.")

        # ---------------- Save Invoice ----------------

        filename = f"{row[0]}.docx"

        doc.save(filename)

        with open(filename, "rb") as file:

            st.download_button(
                "📥 Download Invoice",
                data=file,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # ---------------- Delete Temporary PDF Images ----------------

        for temp_file in os.listdir():

            if temp_file.startswith("temp_page_") and temp_file.endswith(".png"):

                os.remove(temp_file)

# ==========================================
# CLOSE DATABASE
# ==========================================

conn.close()